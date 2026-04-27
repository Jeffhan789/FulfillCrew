from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "云端多智能体智能电商系统_项目构建思路说明.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size=None, bold=None, color=None, latin="Aptos", east_asia="Microsoft YaHei"):
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    return run


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, bold=bold, color=color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(
            run,
            color=RGBColor(31, 78, 121) if level == 1 else RGBColor(44, 88, 67),
            latin="Aptos Display",
        )
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run)


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run)


def add_table(doc, headers, rows, widths=None):
    header_line = " / ".join(headers)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(header_line)
    set_run_font(run, size=9.5, bold=True, color=RGBColor(31, 78, 121))

    for row in rows:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        for idx, value in enumerate(row):
            if idx == 0:
                label = paragraph.add_run(f"{value}：")
                set_run_font(label, bold=True, color=RGBColor(44, 88, 67))
            else:
                text = "；" if idx > 1 else ""
                run = paragraph.add_run(f"{text}{value}")
                set_run_font(run)
    doc.add_paragraph()
    return None


def add_callout(doc, title, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    title_run = paragraph.add_run(f"【{title}】")
    set_run_font(title_run, bold=True, color=RGBColor(44, 88, 67))
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(line)
        set_run_font(run)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("云端多智能体智能电商系统")
    set_run_font(run, size=24, bold=True, color=RGBColor(31, 78, 121), latin="Aptos Display")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("项目构建思路说明文档")
    set_run_font(run, size=14, color=RGBColor(74, 85, 92))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    run = meta.add_run("GitHub 项目名：cloud-multi-agent-ecommerce-intelligence-system")
    set_run_font(run, size=10, color=RGBColor(90, 90, 90))

    add_callout(
        doc,
        "一句话定位",
        [
            "本项目把 COMP315 的两个 coursework 从“数据清洗 + React 电商前端”升级为一个完整的 AI + Multi-Agent 电商系统。",
            "它的价值不在于从零堆功能，而在于把已有课程成果继续工程化、智能化、系统化，形成可展示、可解释、可迭代的 portfolio project。",
        ],
    )

    add_heading(doc, "1. 项目为什么这样构建", 1)
    add_body(
        doc,
        "这个项目的核心策略是“复用已有成果，扩展系统深度”。COMP315 CA1 已经提供商品数据清洗能力，CA2 已经提供电商前端雏形，因此第一版不需要重新发明一个电商网站，而是把这两个作业作为工程底座。",
    )
    add_body(
        doc,
        "在此基础上，项目向两个方向升级：第一，加入后端 API、订单流程和部署结构，让它像真实系统一样运行；第二，加入 Multi-Agent 决策层和 Neural Network 预测层，让它体现 COMP310 与 ELEC320 的课程技术。",
    )

    add_heading(doc, "2. 三门课的融合逻辑", 1)
    add_table(
        doc,
        ["课程", "在项目中的角色", "具体体现"],
        [
            ["COMP315", "云端电商工程底座", "CA1 数据清洗、CA2 React 前端、FastAPI 后端、Docker 部署、API 系统设计"],
            ["COMP310", "多智能体订单决策层", "Order Agent、Inventory Agent、Coordinator Agent、Warehouse Agents、Contract Net bidding"],
            ["ELEC320", "神经网络智能预测层", "Demand Prediction、Fraud Detection、Category Classification、training / online inference 思路"],
        ],
        widths=[2.6, 4.2, 9.0],
    )

    add_heading(doc, "3. 从 CA1 到数据处理层", 1)
    add_body(
        doc,
        "CA1 的 data cleaning 不再只是一个独立作业，而是系统入口的数据 pipeline。所有商品进入前端、后端、Agent 和 ML 模块之前，都先被清洗成稳定字段。",
    )
    add_bullet(doc, "输入：raw product JSON，可能包含空名称、非法价格、错误库存、缺失图片链接等问题。")
    add_bullet(doc, "处理：clean_name、clean_price、clean_category、clean_type、clean_quantity、clean_rating、clean_image_link。")
    add_bullet(doc, "输出：cleaned_products.json，成为前端展示、后端 API、库存判断和模型输入的共同数据源。")

    add_heading(doc, "4. 从 CA2 到前端展示层", 1)
    add_body(
        doc,
        "CA2 的 React E-Commerce Website 是用户可见的第一层。它保留原有搜索、排序、库存筛选、basket 管理和总价计算，再新增 checkout 与 Agent 决策展示。",
    )
    add_bullet(doc, "用户在前端浏览商品，并通过搜索、排序和 in-stock 过滤找到可购买商品。")
    add_bullet(doc, "用户把商品加入 basket 后点击 checkout。")
    add_bullet(doc, "前端调用后端订单 API，并显示订单状态、风险评分、仓库分配、需求预测和 Agent decision log。")

    add_heading(doc, "5. 后端与 Multi-Agent 层", 1)
    add_body(
        doc,
        "后端使用 FastAPI 作为系统中枢。前端不直接处理业务决策，而是把订单提交给后端，由 Agent 层完成库存检查、风险判断、仓库竞价和需求预测。",
    )
    add_table(
        doc,
        ["Agent", "职责", "输出结果"],
        [
            ["Order Agent", "接收 checkout 请求，创建订单流程", "order_status、订单日志"],
            ["Fraud Detection Agent", "调用异常订单检测模型，计算风险分数", "risk_score、approved / review_required"],
            ["Inventory Agent", "检查库存并在订单通过后扣减库存", "stock_available、remaining quantity"],
            ["Coordinator Agent", "向多个 Warehouse Agent 请求 bids 并选择最优仓库", "selected_warehouse、bid log"],
            ["Warehouse Agents", "根据库存、距离、工作负载和处理速度提交 bid", "warehouse bid"],
            ["Demand Prediction Agent", "调用需求预测模型判断未来销量", "predicted_demand、restock_recommendation"],
        ],
        widths=[3.5, 7.2, 5.1],
    )

    add_heading(doc, "6. Neural Network 层的设计思路", 1)
    add_body(
        doc,
        "第一版 MVP 使用轻量、确定性的模型接口来模拟 MLP / 分类器的在线推理过程。这样做的好处是项目可以先完整跑起来，后续再替换成真正训练好的 PyTorch、TensorFlow 或 scikit-learn 模型。",
    )
    add_bullet(doc, "Demand Prediction：输入 price、quantity、rating、category、type 等特征，输出 next 7-day demand。")
    add_bullet(doc, "Fraud Detection：输入 order_total、number_of_items、average_item_price、is_new_user、shipping_distance，输出 0 到 1 的 risk_score。")
    add_bullet(doc, "Product Category Classifier：用商品名称关键词做第一版分类器，后续可升级为 bag-of-words、embedding 或 MLP classifier。")

    add_heading(doc, "7. MVP 完整运行流程", 1)
    steps = [
        "CA1 风格清洗脚本生成 cleaned_products.json。",
        "FastAPI 后端读取 cleaned product data。",
        "React 前端调用 /products API 展示商品列表。",
        "用户搜索、排序、筛选 in-stock 商品，并加入 basket。",
        "用户点击 checkout，前端 POST 到 /orders。",
        "Order Agent 接收订单请求。",
        "Fraud Detection Agent 计算风险分数。",
        "Inventory Agent 检查商品库存。",
        "Coordinator Agent 向 Warehouse A / B / C 请求 bids。",
        "Warehouse Agents 根据负载、距离和库存提交 bids。",
        "Coordinator Agent 选择最优 warehouse。",
        "Demand Prediction Agent 预测未来 7 天需求。",
        "前端显示订单结果和完整 Agent decision log。",
    ]
    for step in steps:
        add_number(doc, step)

    add_heading(doc, "8. 当前项目文件结构如何理解", 1)
    add_table(
        doc,
        ["目录", "作用"],
        [
            ["data_cleaning/", "对应 COMP315 CA1，负责商品 JSON 清洗与 cleaned data 输出。"],
            ["frontend/", "对应 COMP315 CA2，负责商品展示、basket、checkout 和 Agent 结果展示。"],
            ["backend/", "负责 FastAPI、产品 API、订单 API、Agent 编排和业务服务。"],
            ["backend/agents/", "实现 COMP310 的 Multi-Agent System 层。"],
            ["ml_models/", "实现 ELEC320 的需求预测、欺诈检测和分类模型接口。"],
            ["docs/", "存放课程映射、系统设计、Agent 设计、模型设计和本说明文档。"],
            ["tests/", "存放数据清洗与 Agent 流程的基础测试。"],
        ],
        widths=[4.0, 11.8],
    )

    add_heading(doc, "9. 下一步开发路线", 1)
    add_body(doc, "建议按“先可演示，再增强智能，再部署上线”的顺序推进。")
    add_table(
        doc,
        ["阶段", "目标", "完成标准"],
        [
            ["第 1 阶段", "把 CA1/CA2 原始代码迁入当前项目", "原作业功能不丢失，前端能读取 cleaned JSON 或 backend API。"],
            ["第 2 阶段", "完善 checkout 与订单 API", "用户下单后可看到订单状态、风险分数、仓库选择和日志。"],
            ["第 3 阶段", "替换真实 MLP 模型", "保存训练脚本、模型评估指标、train/test split 和预测接口。"],
            ["第 4 阶段", "加入数据库", "用 SQLite 或 PostgreSQL 保存商品、订单、库存和 Agent logs。"],
            ["第 5 阶段", "Docker / 云端部署", "docker-compose 一键启动，README 包含部署说明和演示截图。"],
        ],
        widths=[2.4, 5.8, 7.6],
    )

    add_heading(doc, "10. 简历与 README 表述", 1)
    add_callout(
        doc,
        "英文项目描述",
        [
            "Extended a coursework-based React e-commerce website into an AI-powered multi-agent fulfilment system by adding backend APIs, neural network prediction modules, autonomous agent coordination and containerised deployment.",
        ],
    )
    add_callout(
        doc,
        "中文项目描述",
        [
            "在 COMP315 课程作业的 JavaScript 数据清洗流程和 React 电商网站基础上，扩展出一个 AI 驱动的多智能体订单履约系统，加入后端 API、神经网络预测模块、Agent 协作决策和 Docker 部署结构。",
        ],
    )

    add_heading(doc, "11. 你应该如何向别人解释这个项目", 1)
    add_body(
        doc,
        "最清晰的讲法不是说“我做了很多技术”，而是说“我把三个课程中的技术放进了一个统一业务流程”。COMP315 解决系统如何运行，COMP310 解决订单如何由多个 Agent 协作决策，ELEC320 解决系统如何做智能预测。",
    )
    add_body(
        doc,
        "展示时可以按一次 checkout 演示：用户提交 basket 后，系统返回 risk score、warehouse bids、selected warehouse、predicted demand 和 restock recommendation。这个链路能同时证明前端、后端、Agent、ML 和部署思路。",
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Cloud-Based Multi-Agent E-Commerce Intelligence System")
    set_run_font(run, size=8, color=RGBColor(120, 120, 120))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
